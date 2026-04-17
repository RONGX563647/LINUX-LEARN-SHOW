#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将docx文件转换为markdown格式
"""

import os
import glob
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def convert_paragraph_to_md(paragraph):
    """将段落转换为markdown格式"""
    text = paragraph.text.strip()
    if not text:
        return ""
    
    # 处理标题
    style_name = paragraph.style.name if paragraph.style else ""
    if "Heading 1" in style_name or "标题 1" in style_name:
        return f"# {text}\n"
    elif "Heading 2" in style_name or "标题 2" in style_name:
        return f"## {text}\n"
    elif "Heading 3" in style_name or "标题 3" in style_name:
        return f"### {text}\n"
    elif "Heading 4" in style_name or "标题 4" in style_name:
        return f"#### {text}\n"
    elif "Heading 5" in style_name or "标题 5" in style_name:
        return f"##### {text}\n"
    elif "Heading 6" in style_name or "标题 6" in style_name:
        return f"###### {text}\n"
    
    # 处理列表
    if paragraph.style and "List" in paragraph.style.name:
        # 检查是否是有序列表
        if hasattr(paragraph, '_element'):
            numPr = paragraph._element.pPr
            if numPr is not None:
                numPr_elem = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr_elem is not None:
                    ilvl = numPr_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                    level = int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if ilvl is not None else 0
                    indent = "  " * level
                    return f"{indent}- {text}\n"
        return f"- {text}\n"
    
    # 处理普通段落
    # 检查是否是粗体
    is_bold = False
    for run in paragraph.runs:
        if run.bold:
            is_bold = True
            break
    
    if is_bold and len(text) < 100:  # 短粗体文本可能是小标题
        return f"**{text}**\n"
    
    return f"{text}\n"


def convert_table_to_md(table):
    """将表格转换为markdown格式"""
    if not table.rows:
        return ""
    
    md_table = []
    
    # 获取表格数据
    rows_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ')
            row_data.append(cell_text)
        rows_data.append(row_data)
    
    if not rows_data:
        return ""
    
    # 确定列数
    max_cols = max(len(row) for row in rows_data)
    
    # 创建表头
    header = rows_data[0] if rows_data else []
    while len(header) < max_cols:
        header.append("")
    
    md_table.append("| " + " | ".join(header) + " |")
    md_table.append("| " + " | ".join(["---"] * max_cols) + " |")
    
    # 添加数据行
    for row in rows_data[1:]:
        while len(row) < max_cols:
            row.append("")
        md_table.append("| " + " | ".join(row) + " |")
    
    return "\n".join(md_table) + "\n"


def convert_docx_to_md(docx_path):
    """将docx文件转换为markdown格式"""
    doc = Document(docx_path)
    md_content = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            for paragraph in doc.paragraphs:
                if paragraph._element == element:
                    md_line = convert_paragraph_to_md(paragraph)
                    if md_line:
                        md_content.append(md_line)
                    break
        # 处理表格
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    md_table = convert_table_to_md(table)
                    if md_table:
                        md_content.append("\n" + md_table + "\n")
                    break
    
    return "\n".join(md_content)


def main():
    """主函数"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    docx_files = glob.glob(os.path.join(current_dir, "*.docx"))
    
    # 过滤掉临时文件
    docx_files = [f for f in docx_files if not os.path.basename(f).startswith("~$")]
    
    if not docx_files:
        print("未找到docx文件")
        return
    
    print(f"找到 {len(docx_files)} 个docx文件")
    
    for docx_file in docx_files:
        try:
            print(f"\n正在转换: {os.path.basename(docx_file)}")
            md_content = convert_docx_to_md(docx_file)
            
            # 生成输出文件名
            md_file = os.path.splitext(docx_file)[0] + ".md"
            
            # 写入markdown文件
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            print(f"转换完成: {os.path.basename(md_file)}")
        except Exception as e:
            print(f"转换失败 {os.path.basename(docx_file)}: {str(e)}")
    
    print("\n所有文件转换完成！")


if __name__ == "__main__":
    main()